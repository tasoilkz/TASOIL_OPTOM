import os
from PIL import Image as PILImage
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

GLOVES_PHOTOS_DIR = "перчатки" 

def create_gloves_document():
    doc = docx.Document()
    
    title = doc.add_heading('Прайс-лист: Перчатки и средства защиты оптом', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    items = [
        {
            "name": "Перчатки х/б с ПВХ (стандарт)",
            "desc": "В мешке 780 шт.\nЦена с НДС: 90 тг × 780 = 70 200 тг",
            "img_filename": "glove_1.jpg"
        },
        {
            "name": "Перчатки х/б полосатые (цветные)",
            "desc": "В мешке 600 шт.\nЦена с НДС: 142 тг × 600 = 85 200 тг",
            "img_filename": "glove_2.jpg"
        },
        {
            "name": "Перчатки с латексным покрытием (300#)",
            "desc": "В мешке 720 шт.\nЦена с НДС: 180 тг × 720 = 129 600 тг",
            "img_filename": "glove_3.jpg"
        },
        {
            "name": "Комбинезон защитный",
            "desc": "В мешке 100 шт.\nЦена с НДС: 1 100 тг × 100 = 110 000 тг",
            "img_filename": "suit.jpg"
        },
        {
            "name": "Жилет сигнальный светоотражающий",
            "desc": "Цена за 1 шт: 800 тг с НДС",
            "img_filename": "vest.jpg"
        },
        {
            "name": "Полотно / Ткань в рулонах",
            "desc": "Цена за рулон: 21 000 тг (1.40 / 70м) с НДС",
            "img_filename": "roll.jpg"
        },
        {
            "name": "Перчатки зимние утепленные (-40°C)",
            "desc": "В мешке 420 пар\nЦена с НДС: 880 тг × 420 = 369 600 тг",
            "img_filename": "winter_glove.jpg"
        }
    ]

    for item in items:
        p_name = doc.add_paragraph()
        run_name = p_name.add_run(item["name"])
        run_name.bold = True
        run_name.font.size = Pt(14)
        
        p_desc = doc.add_paragraph(item["desc"])
        p_desc.paragraph_format.space_after = Pt(10)
        
        img_path = os.path.join(GLOVES_PHOTOS_DIR, item["img_filename"])
        
        if os.path.exists(img_path):
            try:
                img = PILImage.open(img_path)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                
                img.thumbnail((800, 800))
                
                os.makedirs("temp_images", exist_ok=True)
                opt_path = os.path.join("temp_images", item["img_filename"])
                img.save(opt_path, "JPEG", quality=85)
                
                doc.add_picture(opt_path, width=Inches(2.5))
            except Exception as e:
                print(f"Ошибка обработки фото {item['img_filename']}: {e}")
        else:
            print(f"Файл не найден: {img_path}")

        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(10)
        doc.add_paragraph("--------------------------------------------------")

    os.makedirs("files", exist_ok=True)
    file_path = "files/gloves_price.docx"
    doc.save(file_path)
    print(f"Файл успешно создан: {file_path}")

if __name__ == "__main__":
    create_gloves_document()